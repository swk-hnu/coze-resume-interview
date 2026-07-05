from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "ppt_summary_source.md"
OUTPUT = Path(r"C:\Users\Lenovo\Desktop\workflow_frontend_summary_fixed.docx")


def set_run_font(run, size=None, bold=False, italic=False):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
      run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(text)
    set_run_font(run, size=16, bold=True)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(text)
    set_run_font(run, italic=True)


def add_paragraph(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def build_doc():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line:
            i += 1
            continue

        if i == 0 and line.startswith("# "):
            add_title(doc, line[2:].strip())
            i += 1
            if i < len(lines) and lines[i].strip():
                add_subtitle(doc, lines[i].strip())
                i += 1
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("- "):
            add_paragraph(doc, line[2:].strip(), style="List Bullet")
        elif any(line.startswith(f"{n}. ") for n in range(1, 10)):
            content = line.split(". ", 1)[1].strip()
            add_paragraph(doc, content, style="List Number")
        else:
            add_paragraph(doc, line.strip())

        i += 1

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_doc()
    print(path)
